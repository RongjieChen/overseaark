#!/usr/bin/env python3
"""Build the current OverseaArk PRD v2.0 DOCX from markdown.

The script intentionally uses python-docx and Word-native styles/fields rather
than rendering a fake visual document. It is scoped to the PRD artifact only.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "PRD-v2.0.md"
TARGET = ROOT / "docs" / "出海方舟OverseaArk-PRD-v2.0.docx"

# Use one broad Unicode face for every script slot so Word, Quick Look, and the
# LibreOffice render gate produce the same mixed Chinese/Latin layout.
FONT_EAST_ASIA = "Arial Unicode MS"
FONT_FALLBACK = "Arial Unicode MS"
FONT_LATIN = "Arial Unicode MS"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None):
    font = run.font
    font.name = FONT_LATIN
    if size:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if color:
        font.color.rgb = RGBColor.from_string(color)
    for script in ("ascii", "hAnsi", "cs"):
        run._element.rPr.rFonts.set(qn(f"w:{script}"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    for theme_attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        run._element.rPr.rFonts.attrib.pop(qn(f"w:{theme_attr}"), None)


def set_style_font(style, size: int, bold: bool = False, color: str = "111827"):
    font = style.font
    font.name = FONT_LATIN
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = RGBColor.from_string(color)
    for script in ("ascii", "hAnsi", "cs"):
        style.element.rPr.rFonts.set(qn(f"w:{script}"), FONT_LATIN)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    for theme_attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        style.element.rPr.rFonts.attrib.pop(qn(f"w:{theme_attr}"), None)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def create_numbering_instance(document: Document, style_name: str = "List Number") -> int:
    """Create a fresh semantic list instance that restarts at one."""
    style_num_id = int(document.styles[style_name].element.pPr.numPr.numId.val)
    numbering = document.part.numbering_part.element
    base_num = next(
        node
        for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == style_num_id
    )
    abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_num_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return new_num_id


def set_paragraph_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = ensure_child(p_pr, "w:numPr")
    level = ensure_child(num_pr, "w:ilvl")
    level.set(qn("w:val"), "0")
    number = ensure_child(num_pr, "w:numId")
    number.set(qn("w:val"), str(num_id))


def apply_table_geometry(table, widths: list[int]):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = ensure_child(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_ind = ensure_child(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    layout = ensure_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for col_idx, width in enumerate(widths):
        table.columns[col_idx].width = Twips(width)
    for row in table.rows:
        row.height = None
        for col_idx, cell in enumerate(row.cells):
            cell.width = Twips(widths[col_idx])
            set_cell_width(cell, widths[col_idx])
            tc_mar = ensure_child(cell._tc.get_or_add_tcPr(), "w:tcMar")
            for side, margin in CELL_MARGINS_DXA.items():
                node = ensure_child(tc_mar, f"w:{side}")
                node.set(qn("w:w"), str(margin))
                node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D1D5DB")


def add_toc(document: Document, markdown: str):
    p = document.add_paragraph(style="TOC Instruction")
    p.add_run("目录").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for line in markdown.splitlines():
        if not line.startswith("## "):
            continue
        text = line[3:].strip()
        if text == "目录":
            continue
        toc_line = document.add_paragraph()
        toc_line.paragraph_format.space_before = Pt(0)
        toc_line.paragraph_format.space_after = Pt(4)
        set_run_font(toc_line.add_run(text), 10, True, "374151")
    document.add_page_break()


def configure_styles(document: Document):
    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal, 11, False, "111827")
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 23, "0F172A"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        set_style_font(styles[name], size, True, color)
    styles["Title"].paragraph_format.space_before = Pt(0)
    styles["Title"].paragraph_format.space_after = Pt(4)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    set_style_font(styles["List Number"], 11, False, "111827")
    set_style_font(styles["List Bullet"], 11, False, "111827")
    for list_name in ("List Number", "List Bullet"):
        list_format = styles[list_name].paragraph_format
        list_format.left_indent = Inches(0.5)
        list_format.first_line_indent = Inches(-0.25)
        list_format.space_after = Pt(8)
        list_format.line_spacing = 1.167
    set_style_font(styles["Caption"], 8, False, "6B7280")

    if "Memo Masthead" not in styles:
        masthead = styles.add_style("Memo Masthead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        masthead = styles["Memo Masthead"]
    set_style_font(masthead, 9, True, "FFFFFF")
    masthead.paragraph_format.space_after = Pt(0)

    if "TOC Instruction" not in styles:
        toc_style = styles.add_style("TOC Instruction", WD_STYLE_TYPE.PARAGRAPH)
    else:
        toc_style = styles["TOC Instruction"]
    set_style_font(toc_style, 12, True, "0F172A")


def configure_document(document: Document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("出海方舟 OverseaArk  |  PRD v2.0")
    set_run_font(run, 8, True, "4B5563")

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = footer.add_run("当前实施与赛事交付基线  ·  ")
    set_run_font(label, 8, False, "6B7280")
    page_field = OxmlElement("w:fldSimple")
    page_field.set(qn("w:instr"), "PAGE")
    page_run = OxmlElement("w:r")
    page_rpr = OxmlElement("w:rPr")
    page_size = OxmlElement("w:sz")
    page_size.set(qn("w:val"), "16")
    page_color = OxmlElement("w:color")
    page_color.set(qn("w:val"), "6B7280")
    page_rpr.extend((page_size, page_color))
    page_text = OxmlElement("w:t")
    page_text.text = "1"
    page_run.extend((page_rpr, page_text))
    page_field.append(page_run)
    footer._p.append(page_field)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows, i


def add_table(document: Document, rows: list[list[str]]):
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    set_table_borders(table)
    header_row_pr = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "true")
    header_row_pr.append(header_marker)

    available_twips = CONTENT_WIDTH_DXA
    first_col = 1900 if max_cols > 2 else 2600
    remaining = available_twips - first_col
    widths = [first_col] + [max(1200, remaining // max(1, max_cols - 1))] * (max_cols - 1)
    widths[-1] += available_twips - sum(widths)

    for r_idx, row in enumerate(rows):
        row_pr = table.rows[r_idx]._tr.get_or_add_trPr()
        if row_pr.find(qn("w:cantSplit")) is None:
            row_pr.append(OxmlElement("w:cantSplit"))
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            value = row[c_idx] if c_idx < len(row) else ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.line_spacing = 1.10
            run = para.add_run(value)
            set_run_font(run, 8 if max_cols >= 5 else 9, r_idx == 0, "111827")
            if r_idx == 0:
                shade_cell(cell, "E5E7EB")
            elif r_idx % 2 == 0:
                shade_cell(cell, "F9FAFB")
    apply_table_geometry(table, widths)


def add_code_block(document: Document, code: list[str]):
    para = document.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.18)
    para.paragraph_format.right_indent = Inches(0.18)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = 1.0
    p_pr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F4F6")
    p_pr.append(shd)
    for idx, line in enumerate(code):
        if idx:
            para.add_run().add_break()
        run = para.add_run(line)
        run.font.name = "Menlo"
        run.font.size = Pt(8)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FALLBACK)


def add_cover(document: Document):
    p = document.add_paragraph(style="Memo Masthead")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    masthead = p.add_run("NVIDIA DGX Spark 黑客松  ·  PRODUCT REQUIREMENTS DOCUMENT")
    set_run_font(masthead, 10, True, "2E74B5")

    title = document.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("DGX Spark一支不下班的\n本地多模态外贸营销团队")
    set_run_font(r, 23, True, "0F172A")
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = sub.add_run("出海方舟 OverseaArk 产品需求文档 v2.0")
    set_run_font(r, 15, True, "374151")
    tagline = document.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = tagline.add_run("一台 NVIDIA DGX Spark 上本地运行、可观测、可恢复、可导出的多模态外贸营销工作台")
    set_run_font(r, 11, False, "4B5563")
    metadata_rows = [
        ("文档版本", "v2.0"),
        ("文档状态", "当前实施与赛事交付基线"),
        ("目标平台", "DGX Spark / Ubuntu 24.04 / aarch64 / CUDA 13"),
        ("技术基线", "native vLLM + Step1X + Cosmos3-Edge + NeMo"),
        ("更新日期", "2026-07-22"),
    ]
    for label, value in metadata_rows:
        metadata = document.add_paragraph()
        metadata.paragraph_format.space_before = Pt(0)
        metadata.paragraph_format.space_after = Pt(2)
        set_run_font(metadata.add_run(f"{label}："), 10, True, "111827")
        set_run_font(metadata.add_run(value), 10, False, "374151")
    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(12)
    rule.paragraph_format.space_after = Pt(0)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:color"), "2E74B5")
    p_bdr.append(bottom)
    rule._p.get_or_add_pPr().append(p_bdr)
    document.add_page_break()


def add_markdown(document: Document, markdown: str):
    lines = markdown.splitlines()
    i = 0
    in_code = False
    code: list[str] = []
    active_numbering_id: int | None = None
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        if line.startswith("```"):
            active_numbering_id = None
            if in_code:
                add_code_block(document, code)
                code = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code.append(raw)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        if line.strip().startswith("|"):
            active_numbering_id = None
            rows, i = parse_table(lines, i)
            add_table(document, rows)
            continue
        if line.startswith("# "):
            active_numbering_id = None
            if "出海方舟" in line:
                i += 1
                continue
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            active_numbering_id = None
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            active_numbering_id = None
            document.add_heading(line[4:].strip(), level=2)
        elif re.match(r"^\d+\.\s+", line):
            if active_numbering_id is None:
                active_numbering_id = create_numbering_instance(document)
            p = document.add_paragraph(style="List Number")
            set_paragraph_numbering(p, active_numbering_id)
            r = p.add_run(re.sub(r"^\d+\.\s+", "", line))
            set_run_font(r, 11)
        elif line.startswith("- "):
            active_numbering_id = None
            p = document.add_paragraph(style="List Bullet")
            r = p.add_run(line[2:])
            set_run_font(r, 11)
        else:
            active_numbering_id = None
            p = document.add_paragraph()
            r = p.add_run(line)
            set_run_font(r, 11)
            next_idx = i + 1
            while next_idx < len(lines) and not lines[next_idx].strip():
                next_idx += 1
            if next_idx < len(lines) and lines[next_idx].strip().startswith(("|", "```")):
                p.paragraph_format.keep_with_next = True
        i += 1


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    document = Document()
    configure_styles(document)
    configure_document(document)
    add_cover(document)
    add_toc(document, markdown)
    add_markdown(document, markdown)
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    document.save(TARGET)
    print(f"wrote {TARGET}")


if __name__ == "__main__":
    main()
